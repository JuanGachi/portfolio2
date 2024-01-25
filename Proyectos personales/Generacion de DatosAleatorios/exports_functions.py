import json
import csv
from docx import Document
import io

# Función para exportar a JSON
def exportar_json(datos, buffer):
    # Convertir los datos JSON a una cadena y luego a bytes
    json_str = json.dumps(datos, indent=4)
    buffer.write(json_str.encode('utf-8'))

# Función para exportar a CSV
def exportar_csv(datos, buffer):
    # Crear un objeto StringIO para escribir los datos CSV
    csv_buffer = io.StringIO()
    writer = csv.writer(csv_buffer, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
    if datos:
        writer.writerow(datos[0].keys())  # Encabezados de columnas
        for dato in datos:
            writer.writerow(dato.values())
    # Escribir los datos CSV en el buffer principal como bytes
    buffer.write(csv_buffer.getvalue().encode('utf-8'))

# Función para exportar a DOCX
def exportar_docx(datos, buffer):
    doc = Document()
    doc.add_heading('Datos Generados', 0)

    if datos:
        # Añadir una tabla para los datos
        tabla = doc.add_table(rows=1, cols=len(datos[0]))
        hdr_cells = tabla.rows[0].cells

        # Configurar los encabezados de la tabla
        for idx, titulo in enumerate(datos[0].keys()):
            hdr_cells[idx].text = titulo

        # Añadir los datos en la tabla
        for dato in datos:
            row_cells = tabla.add_row().cells
            for idx, valor in enumerate(dato.values()):
                row_cells[idx].text = str(valor)

    # Guardar el documento en un buffer en memoria
    fake_file = io.BytesIO()
    doc.save(fake_file)
    fake_file.seek(0)
    buffer.write(fake_file.read())

# Función para exportar a TXT
def exportar_txt(datos, buffer):
    txt_str = '\n'.join([f"{', '.join(str(valor) for valor in dato.values())}" for dato in datos])
    buffer.write(txt_str.encode('utf-8'))
